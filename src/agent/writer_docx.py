from docx import Document
from .config import now_ist

def write_docx(items: list[dict], path: str, title_prefix: str = "Tech Pulse"):
    doc = Document()
    doc.add_heading(f"{title_prefix} — {now_ist().strftime('%Y-%m-%d %H:00 IST')}", level=1)
    for i, it in enumerate(items, 1):
        doc.add_heading(f"{i}. {it['title']}", level=2)
        meta = it["source"]
        if it.get("published"):
            meta += " • " + it["published"].strftime("%Y-%m-%d %H:%M IST")
        doc.add_paragraph(meta)
        # write the agentic summary here
        doc.add_paragraph(it.get("summary", "") or "")
        # link
        doc.add_paragraph(it["link"])
    doc.save(path)
